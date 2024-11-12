from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class Doc:
    def __init__(self, equipment):
        self.equipment = equipment

    def check_equipment_list(self, equipment_list, equipment = ""):
        unique_first_values = set()

        for name, quantity in equipment_list:
            if name in unique_first_values:
                return False
            else:
                unique_first_values.add(name)
            
            if equipment != "" and equipment == name:
                return True
    
    def validate_equipment_list(self, equipment_list):
        isInvalid = self.check_equipment_list(equipment_list)

        if isInvalid:
            raise Exception("Не должно быть дубликатов!")
        
        for name, quantity in equipment_list:
            equipment_found = self.equipment.findByName(name)

            if not equipment_found:
                self.equipment.create(name)

            if not name:
                raise Exception("Не задано название оборудования")

            if not quantity:
                raise Exception(f"Не задано количество для оборудования {name}")

            try:
                quantity = int(quantity)
            except ValueError:
                raise Exception(f"Количество для оборудования {name} должно быть числом")

    def create_act_document(self, equipment_list, object, licenses):
        self.validate_equipment_list(equipment_list)

        document = Document('./doc/templates/act_template.docx')
        replacements = {
            '{{Customer}}': object['customer']['name'],
            '{{Object}}': object['name'],
            '{{Address}}': object['address'],
            '{{License}}': licenses
        }

        for paragraph in document.paragraphs:
            if '{{Equipment}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{Equipment}}', '')
                    act_equipment_table = self.create_act_equipment_table(document, equipment_list)

                    tbl, p = act_equipment_table._element, paragraph._element
                    p.addnext(tbl)
        
        self.make_replacements(document, replacements)
        
        document.save(f'./created_docs/Акт п.о. {replacements["{{Customer}}"]}.docx')

        return f'Акт п.о. {replacements["{{Customer}}"]}.docx'

    def create_passport_document(self, equipment_list, object):
        self.validate_equipment_list(equipment_list)

        document = Document('./doc/templates/passport_template.docx')

        replacements = {
            '{{Customer}}': object['customer']['name'],
            '{{Object}}': f"{object['name']} {object['address']}",
        }

        passport_table = self.create_passport_table(document, equipment_list, object['name'])
        
        for paragraph in document.paragraphs:
            if '{{Object_table}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{Object_table}}', '')

                tbl, p = passport_table._element, paragraph._element
                p.addnext(tbl)
            
        self.make_replacements(document, replacements)
        
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if  "{{Object}}" in paragraph.text:
                            inline = paragraph.runs

                            for item in inline:
                                if  "{{Object}}" in item.text:
                                    item.text = item.text.replace("{{Object}}", replacements['{{Object}}'])

        document.save(f'./created_docs/Паспорт {replacements["{{Customer}}"]}.docx')

        return f'Паспорт {replacements["{{Customer}}"]}.docx'

    def create_act_equipment_table(self, document, equipment_list):
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Наименование, тип установленных элементов система ПА'
        hdr_cells[1].text = 'Количество установленных элементов система ПА'

        for name, quantity in equipment_list:
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = str(quantity)
        
        self.apply_style_to_table(table)

        return table
    
    def create_passport_table(self, document, equipment_list, obj_name):
        N = len(equipment_list)
        table = document.add_table(rows=N + 1, cols=4)  
        table.style = 'Table Grid'

        headers = [
            'Объект', 
            'Наименование установленных элементов ПА', 
            'Кол-во', 
            'Номер и срок действия документа об оценке соответствия установленных элементов ПА'
        ]

        for i, header in enumerate(headers):
            table.cell(0, i).text = header
            table.cell(0, i).paragraphs[0].runs[0].font.size = Pt(12)
            table.cell(0, i).paragraphs[0].runs[0].bold = True
            

        table.cell(1, 0).text = obj_name
        table.cell(1, 0).paragraphs[0].runs[0].font.size = Pt(12)

        for row in range(2, N + 1):
            table.cell(row, 0).merge(table.cell(1, 0))

        for row in range(N):
            table.cell(row + 1, 1).text = equipment_list[row][0]
            table.cell(row + 1, 1).paragraphs[0].runs[0].font.size = Pt(12)
            table.cell(row + 1, 2).text = str(equipment_list[row][1])
            table.cell(row + 1, 2).paragraphs[0].runs[0].font.size = Pt(12)

        table.cell(1, 3).text = 'Акт первичного обследования б/н от '
        table.cell(1, 3).paragraphs[0].runs[0].font.size = Pt(12)

        for row in range(2, N + 1):
            table.cell(row, 3).merge(table.cell(1, 3))
        
        return table
    
    def make_replacements(self, document, replacements):
        for paragraph in document.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    parts = paragraph.text.split(key)
                    paragraph.clear()  
                    run = paragraph.add_run(parts[0])

                    self.set_font_style(run)

                    run = paragraph.add_run(value)

                    self.set_font_style(run)

                    if len(parts) > 1:
                        run = paragraph.add_run(parts[1])
                        self.set_font_style(run)

    def set_font_style(self, run, font_name='Times New Roman', font_size=14):
        run.font.name = font_name
        run.font.size = Pt(font_size)

    def apply_style_to_table(self, table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
